import zipfile
import re
import io
import os
import xml.etree.ElementTree as ET
from typing import Optional

COMPANIES = {
    '志昌資產管理股份有限公司': {
        'name': '志昌資產管理股份有限公司',
        'person': '連偉策',
        'id': '90634048',
        'contact_label': '電子郵件',
        'contact': 'sevice@zcasset.com.tw',
        'address': '臺北市中山區長安東路2段80號10樓之1',
        'bank': '國泰世華銀行 慶城分行',
        'account_name': '志昌資產管理股份有限公司',
        'account_no': '268035011822',
    },
    '瀚昱開發股份有限公司': {
        'name': '瀚昱開發股份有限公司',
        'person': '錢漢洲',
        'id': '62205204',
        'contact_label': '電子郵件',
        'contact': 'service@hanyudev.com',
        'address': '臺北市中山區松江路50號9樓',
        'bank': '凱基商業銀行 城東分行',
        'account_name': '瀚昱開發股份有限公司',
        'account_no': '60070100038023',
    },
    '毅源開發股份有限公司': {
        'name': '毅源開發股份有限公司',
        'person': '吳品毅',
        'id': '62204330',
        'contact_label': '電子郵件',
        'contact': 'service@yiyuandev.com',
        'address': '臺北市松山區寶清街21號4樓之1',
        'bank': '凱基商業銀行 城東分行',
        'account_name': '毅源開發股份有限公司',
        'account_no': '60070100038078',
    },
}

USPACE = {
    'name': '悠勢科技股份有限公司',
    'person': '宋捷仁',
    'id': '52492792',
    'contact_label': '聯絡電話',
    'contact': '02-7751-8097',
    'address': '臺北市中山區八德路二段232號9樓',
}

# Fixed per-company signature page field lines (label + value)
SIG_FIELDS = {
    '志昌資產管理股份有限公司': [
        '甲方名稱：志昌資產管理股份有限公司',
        '負責人：連偉策',
        '統一編號：90634048',
        '電子信箱：sevice@zcasset.com.tw',
        '聯絡地址：臺北市中山區長安東路2段80號10樓之1',
    ],
    '瀚昱開發股份有限公司': [
        '甲方名稱：瀚昱開發股份有限公司',
        '負責人：錢漢洲',
        '統一編號：62205204',
        '地址：臺北市中山區松江路50號9樓',
        '電子信箱：service@hanyudev.com',
    ],
    '毅源開發股份有限公司': [
        '甲方名稱：毅源開發股份有限公司',
        '負責人：吳品毅',
        '統一編號：62204330',
        '地址：臺北市松山區寶清街21號4樓之1',
        '電子信箱：service@yiyuandev.com',
    ],
}

USPACE_SIG_FIELDS = [
    '乙方名稱：悠勢科技股份有限公司',
    '代表人： 宋捷仁',
    '統一編號：52492792',
    '聯絡電話：02-7751-8097',
    '聯絡地址：臺北市中山區八德路二段232號9樓',
]

PIC_DIR = os.path.join(os.path.dirname(__file__), 'pic')

NEW_TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), '自動產生合約範本')

BANK_IMAGES = {
    '瀚昱開發股份有限公司': '瀚昱-凱基城東.jpg',
    '毅源開發股份有限公司': '毅源-凱基城東.jpg',
}

PPR_SIG = (
    '<w:pPr>'
    '<w:spacing w:line="420" w:lineRule="auto"/>'
    '<w:jc w:val="both"/>'
    '<w:rPr>'
    '<w:rFonts w:ascii="思源黑體" w:eastAsia="思源黑體" w:hAnsi="思源黑體" w:cs="思源黑體"/>'
    '<w:color w:val="000000"/>'
    '</w:rPr>'
    '</w:pPr>'
)
RPR_SIG = (
    '<w:rPr>'
    '<w:rFonts w:ascii="思源黑體" w:eastAsia="思源黑體" w:hAnsi="思源黑體" w:cs="思源黑體"/>'
    '<w:color w:val="000000"/>'
    '</w:rPr>'
)


def make_ppr_sig(font: str) -> str:
    return (
        '<w:pPr>'
        '<w:spacing w:line="420" w:lineRule="auto"/>'
        '<w:jc w:val="both"/>'
        '<w:rPr>'
        f'<w:rFonts w:ascii="{font}" w:eastAsia="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
        '<w:color w:val="000000"/>'
        '</w:rPr>'
        '</w:pPr>'
    )


def xml_escape(text: str) -> str:
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


def detect_income_code(plain_text: str) -> str:
    """Detect income code (所得代號) from plain contract text."""
    if re.search(r'\b51L\b', plain_text):
        return '空地租賃(51L)'
    if re.search(r'\b51J\b', plain_text):
        return '建物租賃(51J)'
    if re.search(r'(?:所得代號|代號)[：:]\s*00\b', plain_text):
        return '00發票'
    if re.search(r'\b00\b.*?發票|發票.*?\b00\b', plain_text):
        return '00發票'
    return ''


def update_header_income_code(header_xml: str, income_code: str) -> str:
    """Insert income_code value after 所得代號： label in a header XML string."""
    if not income_code:
        return header_xml
    return re.sub(
        r'(<w:t[^>]*>)([^<]*所得代號[：:])([^<]*)(</w:t>)',
        lambda m: f'{m.group(1)}{m.group(2)}{income_code}{m.group(4)}',
        header_xml,
    )


def _strip_highlights(xml: str) -> str:
    xml = re.sub(r'<w:highlight\b[^/]*/>', '', xml)
    xml = re.sub(r'<w:shd\b[^/]*/>', '', xml, flags=re.DOTALL)
    return xml


def make_rpr(font: str) -> str:
    return (f'<w:rPr>'
            f'<w:rFonts w:ascii="{font}" w:eastAsia="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
            f'<w:color w:val="000000"/>'
            f'</w:rPr>')


def sig_para(text: str, font: str = '思源黑體', size: int = 0) -> str:
    ppr = make_ppr_sig(font)
    rpr = make_rpr(font)
    if size:
        rpr = rpr.replace('</w:rPr>', f'<w:sz w:val="{size * 2}"/><w:szCs w:val="{size * 2}"/></w:rPr>')
    return (
        f'<w:p>{ppr}'
        f'<w:r>{rpr}'
        f'<w:t xml:space="preserve">{xml_escape(text)}</w:t>'
        f'</w:r></w:p>'
    )


def empty_sig_para(font: str = '思源黑體') -> str:
    return f'<w:p>{make_ppr_sig(font)}</w:p>'


def get_plain_text(xml: str) -> str:
    text = re.sub('<[^>]+>', '', xml)
    return re.sub(r'\s+', ' ', text).strip()


def detect_main_font(xml: str) -> str:
    """Find the most common Chinese font in the document."""
    import collections
    fonts = re.findall(r'w:eastAsia="([^"]+)"', xml)
    if not fonts:
        return '思源黑體'
    skip = {'Noto Sans', 'Arial', 'Times New Roman', 'Calibri', 'Tahoma'}
    filtered = [f for f in fonts if f not in skip]
    candidates = filtered if filtered else fonts
    return collections.Counter(candidates).most_common(1)[0][0]


# Old name before company restructuring → canonical current name
_COMPANY_ALIASES = {
    '志昌資產管理有限公司': '志昌資產管理股份有限公司',
}

_COMPANY_SHORT = {
    '志昌': '志昌資產管理股份有限公司',
    '瀚昱': '瀚昱開發股份有限公司',
    '毅源': '毅源開發股份有限公司',
}

def detect_company(plain_text: str) -> Optional[str]:
    # First try full name (fast path)
    for name in COMPANIES:
        if name in plain_text:
            return name
    # Try known old/alias names (e.g. 志昌改名前：志昌資產管理有限公司)
    for alias, canonical in _COMPANY_ALIASES.items():
        if alias in plain_text:
            return canonical
    # Fallback: OOXML may insert spaces mid-name; try compact text
    compact = re.sub(r'\s+', '', plain_text)
    for name in COMPANIES:
        if re.sub(r'\s+', '', name) in compact:
            return name
    for alias, canonical in _COMPANY_ALIASES.items():
        if re.sub(r'\s+', '', alias) in compact:
            return canonical
    # Last resort: match on short distinctive prefix
    for short, full in _COMPANY_SHORT.items():
        if short in plain_text or short in compact:
            return full
    return None


def list_paragraphs(xml: str) -> list:
    """Return list of (start, end, plain_text) for all top-level paragraphs."""
    result = []
    pos = 0
    while True:
        start = xml.find('<w:p', pos)
        if start == -1:
            break
        # Only match <w:p> or <w:p ...>, not <w:pPr>, <w:pStyle>, etc.
        if len(xml) > start + 4 and xml[start + 4] not in (' ', '>'):
            pos = start + 4
            continue
        end = xml.find('</w:p>', start)
        if end == -1:
            break
        end += 6
        para_xml = xml[start:end]
        text = re.sub('<[^>]+>', '', para_xml)
        text = re.sub(r'\s+', ' ', text).strip()
        result.append((start, end, text))
        pos = end
    return result


def get_ppr(para_xml: str) -> str:
    """Extract <w:pPr>...</w:pPr> from a paragraph, or return default."""
    for tag in ('<w:pPr>', '<w:pPr '):
        start = para_xml.find(tag)
        if start != -1:
            end = para_xml.find('</w:pPr>')
            if end != -1:
                return para_xml[start:end + 8]
    return PPR_SIG


def build_header_para(ppr: str, label_value: str, suffix: str) -> str:
    spaces = '                                         '
    return (
        f'<w:p>{ppr}'
        f'<w:r>{RPR_SIG}'
        f'<w:t xml:space="preserve">{xml_escape(label_value)}</w:t>'
        f'</w:r>'
        f'<w:r>{RPR_SIG}'
        f'<w:t xml:space="preserve">{spaces}{xml_escape(suffix)}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )


def detect_sig_label_format(xml: str, sig_idx: int, paragraphs: list) -> str:
    """Detect what party label style is used in the signature section."""
    snippet = ' '.join(p[2] for p in paragraphs[sig_idx:min(sig_idx + 15, len(paragraphs))])
    if '甲方（蓋章）' in snippet or '乙方（蓋章）' in snippet:
        return 'stamp'
    if re.search(r'甲\s*方[：:]', snippet) and re.search(r'乙\s*方[：:]', snippet):
        return 'spaced'
    if '甲方名稱' in snippet or '乙方名稱' in snippet:
        return 'name'
    return 'stamp'


def build_signature_section(company: dict, font: str = '思源黑體',
                             label_fmt: str = 'stamp', sig_title: str = '立契約書人') -> str:
    c = company
    u = USPACE

    if label_fmt == 'spaced':
        jia = f'甲　方：{c["name"]}　（以下簡稱甲方）'
        yi  = f'乙　方：{u["name"]}　（以下簡稱乙方）'
        parts = [
            sig_para(sig_title, font),
            empty_sig_para(),
            sig_para(jia, font),
            sig_para(yi, font),
            empty_sig_para(),
            sig_para('中　華　民　國　　　年　　月　　日', font, size=20),
        ]
    else:  # stamp / name / default
        jia_label = '甲方名稱：' if label_fmt == 'name' else '甲方（蓋章）：'
        yi_label  = '乙方名稱：' if label_fmt == 'name' else '乙方（蓋章）：'
        parts = [
            sig_para(sig_title, font),
            empty_sig_para(),
            sig_para(f'{jia_label}{c["name"]}', font),
            sig_para(f'負責人/出租人： {c["person"]}', font),
            sig_para(f'身分證字號/統編：{c["id"]}', font),
            sig_para(f'{c["contact_label"]}：{c["contact"]}', font),
            sig_para(f'地址：{c["address"]}', font),
            empty_sig_para(),
            sig_para(f'{yi_label}{u["name"]}', font),
            sig_para(f'負責人/承租人： {u["person"]}', font),
            sig_para(f'身分證字號/統編：{u["id"]}', font),
            sig_para(f'{u["contact_label"]}：{u["contact"]}', font),
            sig_para(f'地址：{u["address"]}', font),
            empty_sig_para(),
            sig_para('中　華　民　國　　　年　　月　　日', font, size=20),
        ]
    return ''.join(parts)


def find_sig_section_idx(paragraphs: list, company_name: str) -> int:
    """
    Find the paragraph index where the signature section starts.
    Tries multiple strategies to handle different contract formats.
    Returns -1 if not found.
    """
    n = len(paragraphs)

    # Strategy 1: Find the LAST "立契約書人" / "立補充協議書人" etc.
    # Using LAST occurrence so contracts with both a brief top section and a
    # detailed bottom section (like 補充協議書) use the detailed bottom one.
    SIG_MARKERS = ('立契約書人', '立補充協議書人', '立協議書人', '立合約書人', '立租賃契約書人')
    last_sig = -1
    for i, (_, _, text) in enumerate(paragraphs):
        text_norm = re.sub(r'[　\s]+', '', text)
        if any(m in text_norm for m in SIG_MARKERS):
            last_sig = i
    if last_sig != -1:
        return last_sig

    # Strategy 2: "(簽名頁如後)" → next section is signature page
    for i, (_, _, text) in enumerate(paragraphs):
        if '簽名頁如後' in text:
            # Return next non-empty paragraph
            for j in range(i + 1, n):
                if paragraphs[j][2].strip():
                    return j
            return i + 1

    # Strategy 3: Company name in last 40% of document → scan back for party block start
    start_from = n * 6 // 10
    company_idx = None
    for i in range(start_from, n):
        if company_name in paragraphs[i][2]:
            company_idx = i
            break

    if company_idx is not None:
        # Scan backwards to find the start of the party block
        # Look for a 甲方 label or "下稱" which signals start of party section
        for j in range(company_idx, max(company_idx - 20, start_from - 3) - 1, -1):
            text = paragraphs[j][2].strip()
            if any(kw in text for kw in ['甲方', '出租人', '下稱「甲方', '（甲方）']):
                return j
        # If no label found, go back a few paragraphs from company
        return max(company_idx - 5, start_from)

    # Strategy 4: "中　華　民　國" date line → go back to find party block
    for i in range(n - 1, n * 5 // 10, -1):
        text = paragraphs[i][2]
        if '中' in text and '華' in text and '民' in text and '國' in text and '年' in text:
            # Found date line; scan backwards for party block start
            for j in range(i, max(i - 25, 0) - 1, -1):
                text_j = paragraphs[j][2].strip()
                if any(kw in text_j for kw in ['甲方', '出租人', '立契約書人']):
                    return j
            return max(i - 15, 0)

    return -1


def replace_party_paras_in_place(xml: str, company: dict, sig_idx: int,
                                   paragraphs: list, font: str, label_fmt: str) -> str:
    """
    For top-of-document signatures (only a brief top party section, no bottom section).
    Replaces party VALUES while preserving original labels and suffixes.
    """
    c = company
    u = USPACE
    search_end = min(sig_idx + 25, len(paragraphs))
    to_replace = []

    for abs_i in range(sig_idx, search_end):
        text = paragraphs[abs_i][2].strip()

        # 甲方 paragraph (business owner → company)
        if re.search(r'^(甲[　\s]*方|出租人)[（：:]', text):
            lbl, _, sfx = _split_label_value_suffix(text, c['name'])
            to_replace.append((abs_i, lbl + c['name'] + sfx))

        # 乙方 paragraph with company name (company → 悠勢)
        elif re.search(r'^(乙[　\s]*方|承租人)[（：:]', text) and c['name'] in text:
            lbl, _, sfx = _split_label_value_suffix(text, u['name'])
            to_replace.append((abs_i, lbl + u['name'] + sfx))

    for abs_i, new_text in reversed(to_replace):
        xml = _rebuild_party_para(xml, abs_i, list_paragraphs(xml), new_text, font)

    return xml


def update_bottom_sig_in_place(xml: str, company: dict, sig_idx: int,
                                paragraphs: list, font: str) -> str:
    """
    Rebuild the bottom signature section using fixed per-company field templates.
    Preserves the sig-title paragraph (and any empty paras before the 甲方 block)
    and the date line onward.
    """
    n = len(paragraphs)
    search_end = min(sig_idx + 60, n)
    company_name = company['name']

    # Find the first 甲方-labeled paragraph (start of block to replace)
    jia_start_idx = None
    for i in range(sig_idx, search_end):
        text_norm = re.sub(r'[　\s]+', '', paragraphs[i][2].strip())
        if re.match(r'^甲方', text_norm):
            jia_start_idx = i
            break

    if jia_start_idx is None:
        # fallback: first non-empty paragraph after sig title
        jia_start_idx = sig_idx + 1
        while jia_start_idx < search_end and not paragraphs[jia_start_idx][2].strip():
            jia_start_idx += 1

    # Find the date line (preserve it and everything after)
    date_idx = None
    for i in range(jia_start_idx, search_end):
        text_norm = re.sub(r'[　\s]+', '', paragraphs[i][2])
        if '中華民國' in text_norm and '年' in paragraphs[i][2]:
            date_idx = i
            break

    # Build fixed replacement content from per-company templates
    jia_fields = SIG_FIELDS.get(company_name, [])
    new_chunks = []
    for line in jia_fields:
        new_chunks.append(sig_para(line, font))
    new_chunks.append(empty_sig_para(font))
    for line in USPACE_SIG_FIELDS:
        new_chunks.append(sig_para(line, font))
    new_chunks.append(empty_sig_para(font))
    new_xml = ''.join(new_chunks)

    # Replace from jia_start to just before date line (or end of search area)
    jia_xml_start = paragraphs[jia_start_idx][0]
    if date_idx is not None:
        replace_xml_end = paragraphs[date_idx][0]
    else:
        replace_xml_end = paragraphs[min(search_end - 1, n - 1)][1]

    return xml[:jia_xml_start] + new_xml + xml[replace_xml_end:]


def replace_signature_section(xml: str, company: dict, font: str = '思源黑體') -> str:
    paragraphs = list_paragraphs(xml)
    company_name = company['name']

    sig_idx = find_sig_section_idx(paragraphs, company_name)
    if sig_idx == -1:
        return xml

    n = len(paragraphs)

    if sig_idx < n * 0.20:
        label_fmt = detect_sig_label_format(xml, sig_idx, paragraphs)
        return replace_party_paras_in_place(xml, company, sig_idx, paragraphs, font, label_fmt)

    return update_bottom_sig_in_place(xml, company, sig_idx, paragraphs, font)


def _split_label_value_suffix(text: str, known_value: str = None) -> tuple:
    """
    Split paragraph text into (label, value, suffix).
    label  = everything up to and including the last ：
    value  = known_value if provided, else content between label and suffix
    suffix = trailing （以下簡稱…）or（下稱…）etc.
    """
    colon_pos = max(text.rfind('：'), text.rfind(':'))
    if colon_pos < 0:
        return text, '', ''
    label = text[:colon_pos + 1]
    rest = text[colon_pos + 1:]
    sfx_match = re.search(r'[　\s]*[（(]', rest)
    if sfx_match:
        suffix = rest[sfx_match.start():]
        value = rest[:sfx_match.start()].strip()
    else:
        suffix = ''
        value = rest.strip()
    return label, known_value if known_value is not None else value, suffix


def _rebuild_party_para(xml: str, para_idx: int, paragraphs: list, new_text: str, font: str) -> str:
    """Replace a single party paragraph with new_text, preserving pPr and rPr."""
    paragraphs = list_paragraphs(xml)
    if para_idx >= len(paragraphs):
        return xml
    ps, pe, _ = paragraphs[para_idx]
    para_xml = xml[ps:pe]
    ppr = get_ppr(para_xml)
    # 1. Try run-level rPr (most specific)
    rpr_m = re.search(r'<w:r[\s>].*?<w:rPr>(.*?)</w:rPr>', para_xml, re.DOTALL)
    if rpr_m:
        rpr = f'<w:rPr>{rpr_m.group(1)}</w:rPr>'
    else:
        # 2. Try paragraph-mark rPr from pPr (often carries the same font when runs inherit it)
        ppr_rpr_m = re.search(r'<w:rPr>(.*?)</w:rPr>', para_xml, re.DOTALL)
        rpr = f'<w:rPr>{ppr_rpr_m.group(1)}</w:rPr>' if ppr_rpr_m else make_rpr(font)
    new_p = (f'<w:p>{ppr}'
             f'<w:r>{rpr}'
             f'<w:t xml:space="preserve">{xml_escape(new_text)}</w:t>'
             f'</w:r></w:p>')
    return xml[:ps] + new_p + xml[pe:]


def update_header_parties(xml: str, company_name: str, font: str = '') -> str:
    """
    Update party names in the contract header/preamble.
    Preserves original labels (甲方/乙方/出租人/承租人/etc.) and suffixes.
    """
    paragraphs = list_paragraphs(xml)
    yi_idx = jia_idx = None

    # Patterns allow parenthetical annotations between label and colon,
    # e.g. "乙方（承租人）：" or "甲方（出租人）："
    # Also handle suffix format: "服務提供方：公司名稱（下稱「乙方」）"
    _COLON_PART = r'[^：:\n]{0,30}[：:]'
    yi_pat = re.compile(r'(?:乙[　\s]*方|承租人)' + _COLON_PART)
    yi_suffix_pat = re.compile(r'[「（(]乙[　\s]*方[」）)]')
    jia_pat = re.compile(r'(?:甲[　\s]*方|出租人)' + _COLON_PART)
    jia_suffix_pat = re.compile(r'[「（(]甲[　\s]*方[」）)]')

    # Resolve canonical name (handles old-name aliases)
    canonical_name = COMPANIES.get(company_name, {}).get('name', company_name)
    # All keys that refer to the same canonical company (for searching old contracts)
    name_variants = [k for k, v in COMPANIES.items() if v['name'] == canonical_name]

    # Find 乙方 paragraph containing the company name (search up to 50 paragraphs)
    for i, (_, _, text) in enumerate(paragraphs[:50]):
        if any(n in text for n in name_variants) and (yi_pat.search(text) or yi_suffix_pat.search(text)):
            yi_idx = i
            break

    if yi_idx is None:
        return xml

    # Find 甲方 paragraph: scan up to 20 paragraphs before yi_idx
    for i in range(yi_idx - 1, max(-1, yi_idx - 20), -1):
        _, _, text = paragraphs[i]
        if jia_pat.search(text.strip()) or jia_suffix_pat.search(text.strip()):
            jia_idx = i
            break

    if not font:
        font = detect_main_font(xml)

    # 乙方 paragraph → replace company name with 悠勢, preserve label+suffix
    yi_text = paragraphs[yi_idx][2]
    yi_lbl, _, yi_sfx = _split_label_value_suffix(yi_text, USPACE['name'])
    yi_prefix = yi_lbl + USPACE['name']

    to_update = []
    # 甲方 paragraph → replace business owner with canonical company name, preserve label+suffix
    if jia_idx is not None:
        jia_text = paragraphs[jia_idx][2]
        jia_lbl, _, jia_sfx = _split_label_value_suffix(jia_text, canonical_name)
        jia_prefix = jia_lbl + canonical_name
        # Align suffixes: pad shorter prefix with full-width spaces
        max_pre = max(len(yi_prefix), len(jia_prefix))
        yi_prefix  = yi_prefix  + '　' * (max_pre - len(yi_prefix))
        jia_prefix = jia_prefix + '　' * (max_pre - len(jia_prefix))
        to_update.append((jia_idx, jia_prefix + jia_sfx))

    to_update.append((yi_idx, yi_prefix + yi_sfx))

    # Apply later index first
    for idx, new_text in sorted(to_update, key=lambda x: x[0], reverse=True):
        xml = _rebuild_party_para(xml, idx, list_paragraphs(xml), new_text, font)

    return xml


def fill_bank_account(xml: str, company: dict) -> str:
    """
    Fill 甲方指定匯款銀行帳戶 with company bank info.
    Handles both table format (租金) and paragraph format (分潤).
    """
    if 'bank' not in company:
        return xml

    marker = xml.find('甲方指定匯款銀行帳戶')
    if marker == -1:
        return xml

    values = [company['bank'], company['account_name'], company['account_no']]
    tbl_start = xml.find('<w:tbl>', marker)

    # --- 租金版：表格格式 ---
    if tbl_start != -1 and tbl_start < marker + 500:
        tbl_end = xml.find('</w:tbl>', tbl_start) + 8
        tbl_xml = xml[tbl_start:tbl_end]
        new_tbl = tbl_xml
        tr_pos = 0
        row_idx = 0
        rpr = ('<w:rPr>'
               '<w:rFonts w:ascii="Microsoft JhengHei UI" w:eastAsia="Microsoft JhengHei UI"'
               ' w:hAnsi="Microsoft JhengHei UI" w:cs="思源黑體"/>'
               '<w:color w:val="000000"/>'
               '</w:rPr>')
        while row_idx < len(values):
            tr_start = new_tbl.find('<w:tr', tr_pos)
            if tr_start == -1:
                break
            tr_end = new_tbl.find('</w:tr>', tr_start) + 7
            tr_xml = new_tbl[tr_start:tr_end]
            tc1_end = tr_xml.find('</w:tc>') + 7
            tc2_start = tr_xml.find('<w:tc', tc1_end)
            tc2_end = tr_xml.find('</w:tc>', tc2_start) + 7
            if tc2_start == -1:
                tr_pos = tr_end
                row_idx += 1
                continue
            tc2_xml = tr_xml[tc2_start:tc2_end]
            p_start = tc2_xml.find('<w:p')
            p_end = tc2_xml.find('</w:p>', p_start) + 6
            ppr = get_ppr(tc2_xml[p_start:p_end])
            new_p = (f'<w:p>{ppr}<w:r>{rpr}'
                     f'<w:t xml:space="preserve">{xml_escape(values[row_idx])}</w:t>'
                     f'</w:r></w:p>')
            new_tc2 = tc2_xml[:p_start] + new_p + tc2_xml[p_end:]
            new_tr = tr_xml[:tc2_start] + new_tc2 + tr_xml[tc2_end:]
            new_tbl = new_tbl[:tr_start] + new_tr + new_tbl[tr_end:]
            tr_pos = tr_start + len(new_tr)
            row_idx += 1
        return xml[:tbl_start] + new_tbl + xml[tbl_end:]

    # --- 分潤版：段落格式 ---
    # Replace the ENTIRE paragraph so any pre-filled 業主 value is removed.
    # Use list_paragraphs to get exact <w:p> boundaries (avoids hitting <w:pPr>).
    label_value_map = [
        ('銀行名稱（含分行）', '銀行名稱（含分行）：', company['bank']),
        ('帳戶名稱：',         '帳戶名稱：',           company['account_name']),
        ('帳戶號碼：',         '帳戶號碼：',           company['account_no']),
    ]
    search_from = marker
    for search_key, label_text, value in label_value_map:
        # Find the paragraph containing this label using list_paragraphs
        paragraphs = list_paragraphs(xml)
        target = None
        for ps, pe, text in paragraphs:
            if search_key in text and ps >= search_from:
                target = (ps, pe, xml[ps:pe])
                break
        if target is None:
            continue
        p_start, p_end, p_xml = target
        ppr = get_ppr(p_xml)
        # Get run properties from the first <w:r> in the paragraph
        rpr_match = re.search(r'<w:r[\s>].*?<w:rPr>(.*?)</w:rPr>', p_xml, re.DOTALL)
        if rpr_match:
            rpr = f'<w:rPr>{rpr_match.group(1)}</w:rPr>'
        else:
            rpr = RPR_SIG
        # Build clean paragraph: label run + value run (old values discarded)
        new_p = (f'<w:p>{ppr}'
                 f'<w:r>{rpr}<w:t xml:space="preserve">{xml_escape(label_text)}</w:t></w:r>'
                 f'<w:r>{rpr}<w:t xml:space="preserve">{xml_escape(value)}</w:t></w:r>'
                 f'</w:p>')
        xml = xml[:p_start] + new_p + xml[p_end:]
        search_from = p_start + len(new_p)

    return xml


def _fix_income_code_for_uspace(xml: str) -> str:
    """Force income code checkboxes to 法人 00發票 (all others unchecked)."""
    paragraphs = list_paragraphs(xml)
    replacements = []
    for ps, pe, text in paragraphs:
        if '■' not in text and '□' not in text:
            continue
        para_xml = xml[ps:pe]
        if '個人' in text and ('51L' in text or '51J' in text):
            new_para = para_xml.replace('■', '□')
        elif '法人' in text and '00' in text:
            new_para = para_xml.replace('□', '■')
        elif '管委會' in text:
            new_para = para_xml.replace('■', '□')
        else:
            continue
        if new_para != para_xml:
            replacements.append((ps, pe, new_para))
    for ps, pe, new_para in sorted(replacements, key=lambda x: x[0], reverse=True):
        xml = xml[:ps] + new_para + xml[pe:]
    return xml


def _append_bank_image(docx_bytes: bytes, img_data: bytes, img_filename: str) -> bytes:
    """Append a new page with the bank account image using python-docx."""
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(io.BytesIO(docx_bytes))

    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(io.BytesIO(img_data), width=Inches(6))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_header_plain(plain: str, fields: dict) -> None:
    """Extract header field values from plain text.

    Header cells are concatenated without whitespace in plain text, so each
    pattern uses a lookahead to stop before the next label.  Values that still
    contain '：' (a label character) are discarded as blank-template noise.
    """
    def _valid(val: str, max_len: int = 20) -> bool:
        return bool(val) and '：' not in val and ':' not in val and len(val) <= max_len

    m = re.search(r'負責業務[：:]\s*(.+?)(?=建物編號|分潤|租賃|$)', plain)
    if m and 'sales' not in fields:
        val = m.group(1).strip()
        if _valid(val, max_len=10):
            fields['sales'] = val

    m = re.search(r'建物編號[：:]\s*([A-Za-z]\d+)', plain)
    if m and 'building_id' not in fields:
        fields['building_id'] = m.group(1)

    m = re.search(r'建物名稱[：:]\s*(.+?)(?=建物綁定|所得代號|$)', plain)
    if m and 'building_name' not in fields:
        val = m.group(1).strip()
        if _valid(val, max_len=20):
            fields['building_name'] = val

    m = re.search(r'建物綁定電話[：:]\s*([\d\-]+)', plain)
    if m and 'building_phone' not in fields:
        fields['building_phone'] = m.group(1)

    # 稅前/稅後 — just the suffix, prefix is always 分潤 in 合作協議書
    m = re.search(r'(?:分潤|租賃)稅(前|後)', plain)
    if m and 'tax_suffix' not in fields:
        fields['tax_suffix'] = m.group(1)


def _extract_header_fields(docx_bytes: bytes) -> dict:
    """Extract sales/building fields from old contract header XML (and body fallback)."""
    fields = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        for name in sorted(z.namelist()):
            if re.match(r'word/header\d+\.xml', name):
                plain = get_plain_text(z.read(name).decode('utf-8'))
                _parse_header_plain(plain, fields)
        if len(fields) < 2:
            plain = get_plain_text(z.read('word/document.xml').decode('utf-8'))[:600]
            _parse_header_plain(plain, fields)
    return fields

_INCOME_CODE_HEADER = '一般法人(00發票)'

V2_USPACE_TEMPLATE = os.path.join(
    NEW_TEMPLATES_DIR, 'NEW',
    '第三方-悠勢_停車場系統管理與技術服務合作協議書_公版_7.16.docx',
)

_V77_MODE_TITLE = re.compile(r'模式\s*([ABCＡＢＣ])\s*[：:]')

# 模式 C 區塊的結束錨點（模式 C 之後的下一個章節標題）
_V77_C_END_ANCHORS = ('固定收益及保底收益之付款方式', '押金', '分潤結算規則')

# 各欄位抽取失敗時填入的空白底線
_V2_BLANKS = {
    'address':                '　　　　　　　　　　',
    'spots':                  '____',
    'amount':                 '________',
    'pay_freq':               '月結',
    'pay_period':             '一（1）',
    'party_a_percent':        '__',
    'party_b_percent':        '__',
    'min_guarantee':          '________',
    'excess_threshold':       '________',
    'excess_party_a_percent': '__',
    'excess_party_b_percent': '__',
    'start_date':             '____年__月__日',
    'end_date':               '____年__月__日',
}


def is_v77_owner_contract(plain_text: str) -> bool:
    """7.7 業主版公版特徵：內文有「模式 A／B／C：」標題。舊格式合約沒有。"""
    return bool(_V77_MODE_TITLE.search(plain_text))


def _find_mode_blocks(plain: str) -> dict:
    """Split plain text into per-mode blocks: {'a': (start, end, checked), ...}"""
    titles = []
    seen = set()
    for m in _V77_MODE_TITLE.finditer(plain):
        letter = {'Ａ': 'A', 'Ｂ': 'B', 'Ｃ': 'C'}.get(m.group(1), m.group(1)).lower()
        if letter not in seen:
            seen.add(letter)
            titles.append((m.start(), letter))
    titles.sort()

    blocks = {}
    for i, (pos, letter) in enumerate(titles):
        if i + 1 < len(titles):
            end = titles[i + 1][0]
        else:
            # 最後一個模式區塊：找下一個章節錨點
            end = len(plain)
            for anchor in _V77_C_END_ANCHORS:
                a = plain.find(anchor, pos)
                if a != -1:
                    end = min(end, a)
            end = min(end, pos + 800)
        # 勾選判斷：標題前 4 個字元內有 ■/☑ 即為勾選
        prefix = plain[max(0, pos - 4):pos]
        checked = '■' in prefix or '☑' in prefix
        blocks[letter] = (pos, end, checked)
    return blocks


def extract_v77_data(plain: str) -> tuple:
    """Extract contract data from a 7.7 業主版 contract.

    Anchors on stable legal phrasing (新台幣…元整、超額門檻、甲方分得…%)
    so that light custom edits by sales don't break extraction.
    Returns (data dict, warnings list of Chinese field labels).
    """
    data = {}
    warnings = []

    m = re.search(r'坐落於\s*(.+?)\s*，\s*共計\s*([0-9０-９]+)\s*格', plain)
    if m:
        data['address'] = m.group(1).strip()
        data['spots'] = m.group(2)
    else:
        warnings.append('停車場地址／車位數')

    m = re.search(
        r'自民國（下同）\s*(.+?)\s*（下稱「生效日」）\s*起至\s*(.+?)\s*（下稱「到期日」）', plain)
    if not m:
        # 容錯：句型被改過時，直接找「N年N月N日…起至…N年N月N日」
        m = re.search(
            r'自(?:民國)?[^。]{0,15}?([0-9０-９]{2,4}\s*年\s*[0-9０-９]{1,2}\s*月\s*[0-9０-９]{1,2}\s*日)'
            r'[^。]{0,25}?起\s*至\s*[^。]{0,15}?'
            r'([0-9０-９]{2,4}\s*年\s*[0-9０-９]{1,2}\s*月\s*[0-9０-９]{1,2}\s*日)', plain)
    if m:
        data['start_date'] = m.group(1).strip()
        data['end_date'] = m.group(2).strip()
    else:
        warnings.append('合作期間起訖日期')

    blocks = _find_mode_blocks(plain)
    checked = [k for k, v in blocks.items() if v[2]]
    if len(checked) == 1:
        mode = checked[0]
    elif len(blocks) == 1:
        # 業務刪掉沒用到的模式，只留一個 → 就是它
        mode = next(iter(blocks))
    elif len(checked) > 1:
        mode = checked[0]
        warnings.append('收益模式（勾選了多個模式，請確認）')
    else:
        mode = None
        warnings.append('收益模式（A／B／C 無法判斷）')
    data['mode'] = mode

    if mode:
        bs, be, _ = blocks[mode]
        block = plain[bs:be]

        if mode == 'a':
            m = re.search(r'新台幣\s*([\d,，0-9０-９]+)\s*元整', block)
            if m:
                data['amount'] = m.group(1)
            else:
                warnings.append('模式A固定收益金額')
            # 付款頻率／期間：悠勢版範本預設月結，若業主版不是月結必須同步改
            m = re.search(
                r'採\s*([^，。＿_]{1,6}?)\s*制{1,2}\s*，\s*即以每'
                r'\s*([0-9０-９一二三四五六七八九十（）()]{1,8}?)\s*個月為一期', block)
            if m:
                data['pay_freq'] = m.group(1)
                data['pay_months'] = m.group(2)
            else:
                warnings.append('模式A付款頻率／期間（輸出為預設月結，請確認）')

        elif mode == 'b':
            m = re.search(r'甲方分得\s*([\d.０-９]+)\s*[%％]', block)
            m2 = re.search(r'乙方分得\s*([\d.０-９]+)\s*[%％]', block)
            if m and m2:
                data['party_a_percent'] = m.group(1)
                data['party_b_percent'] = m2.group(1)
            else:
                warnings.append('模式B分潤比例')

        elif mode == 'c':
            m = re.search(r'新台幣\s*([\d,，0-9０-９]+)\s*元整', block)
            if m:
                data['min_guarantee'] = m.group(1)
            else:
                warnings.append('模式C保底金額')
            m = re.search(r'超過新台幣\s*([\d,，0-9０-９]+)\s*元', block)
            if m:
                data['excess_threshold'] = m.group(1)
            else:
                warnings.append('模式C超額門檻')
            m = re.search(r'甲方分得\s*([\d.０-９]+)\s*[%％]', block)
            m2 = re.search(r'乙方分得\s*([\d.０-９]+)\s*[%％]', block)
            if m and m2:
                data['excess_party_a_percent'] = m.group(1)
                data['excess_party_b_percent'] = m2.group(1)
            else:
                warnings.append('模式C超額分潤比例')

        # 含稅／未稅：從勾選模式的區塊內抓「停車營收（…）」
        m = re.search(r'停車營收\s*（\s*(含稅|未稅)\s*）', block)
        data['tax_type'] = m.group(1) if m else '未稅'
    else:
        data['tax_type'] = '未稅'

    # 簽約日期：文末「中華民國X年X月X日」（取最後一個有數字的）
    sign_date = None
    for m in re.finditer(
            r'中[　\s]*華[　\s]*民[　\s]*國[　\s]*'
            r'([0-9０-９]{2,3}\s*年\s*[0-9０-９]{1,2}\s*月\s*[0-9０-９]{1,2}\s*日)', plain):
        sign_date = m.group(1)
    if sign_date:
        data['sign_date'] = re.sub(r'\s+', '', sign_date)

    # 客製偵測：展延條款（公版是「不限次數」，被改過就警告）
    m = re.search(r'順延一年\s*[，,]?\s*([^。，,]{1,12})\s*。', plain)
    if m and m.group(1).strip() != '不限次數':
        warnings.append(f'展延條款疑似經客製修改（業主版為「{m.group(1).strip()}」，'
                        f'輸出為公版「不限次數」，請人工確認）')

    return data, warnings


def _convert_v77(docx_bytes: bytes, original_filename: str, company_name: str,
                 plain: str) -> tuple:
    """Render the 悠勢版 7.7 V2 template from data extracted out of a 7.7 業主版 contract."""
    from docxtpl import DocxTemplate
    from generator import _post_process_docx, _override_fonts

    if not os.path.exists(V2_USPACE_TEMPLATE):
        raise ValueError('找不到悠勢版 7.7 V2 範本檔案')

    co = COMPANIES[company_name]
    data, warnings = extract_v77_data(plain)
    header_fields = _extract_header_fields(docx_bytes)

    mode = data.get('mode')
    tax = data.get('tax_type', '未稅')

    ctx = {
        # 甲方＝第三方公司，資料全部查表帶入
        'party_a':         co['name'],
        'owner':           co['person'],
        'id_number':       co['id'],
        'contact_address': co['address'],
        'email':           co['contact'],
        'bank_name':       co['bank'],
        'account_name':    co['account_name'],
        'account_number':  co['account_no'],
        # 乙方悠勢固定寫死在範本內
        'sign_date':   data.get('sign_date', ''),
        'income_code': _INCOME_CODE_HEADER,
        # 內文模式勾選
        'mode_a_check': '■' if mode == 'a' else '□',
        'mode_b_check': '■' if mode == 'b' else '□',
        'mode_c_check': '■' if mode == 'c' else '□',
        # 模式 A 付款頻率／期間（從業主版帶入，預設月結）
        'pay_freq':   data.get('pay_freq', ''),
        'pay_period': data.get('pay_months', ''),
        # header 欄位（從業主版 header 抄過來；7.16 header 稅別為固定文字，無勾選變數）
        'sales':          header_fields.get('sales', ''),
        'building_id':    header_fields.get('building_id', ''),
        'building_name':  header_fields.get('building_name', ''),
        'building_phone': header_fields.get('building_phone', ''),
    }
    for key in ('address', 'spots', 'start_date', 'end_date', 'amount',
                'party_a_percent', 'party_b_percent', 'min_guarantee',
                'excess_threshold', 'excess_party_a_percent', 'excess_party_b_percent'):
        ctx[key] = data.get(key, '')

    # 7.16 新增履約保證金條款，業主合約無此資料，一律提醒人工填寫
    warnings.append('履約保證金（金額、繳交／退還工作日）為 7.16 新增條款，請人工填寫')

    for field, blank in _V2_BLANKS.items():
        if not str(ctx.get(field, '')).strip():
            ctx[field] = blank

    tpl = DocxTemplate(V2_USPACE_TEMPLATE)
    tpl.render(ctx)
    buf = io.BytesIO()
    tpl.save(buf)
    output_bytes = _post_process_docx(buf.getvalue(), data.get('sign_date', ''))
    output_bytes = _override_fonts(output_bytes)

    base = original_filename[:-5] if original_filename.lower().endswith('.docx') else original_filename
    output_filename = f'{base}(對悠勢合約).docx'

    return output_bytes, output_filename, warnings


def validate_xml(xml_str: str) -> None:
    """Raise ValueError if XML is not well-formed."""
    try:
        # Register common OOXML namespaces to avoid parse errors
        ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError as e:
        raise ValueError(f'合約 XML 格式有誤，無法轉換此檔案（{e}）')


def _convert_party_swap(docx_bytes: bytes, original_filename: str,
                        company_name: str) -> tuple:
    """非公版合約（增補合約等）：內容完全不動，只替換甲乙方。
    甲方 → 偵測到的第三方公司，乙方 → 悠勢科技。
    同時填入甲方銀行帳戶、所得代號固定法人 00發票。"""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read('word/document.xml').decode('utf-8')
        try:
            styles_xml = z.read('word/styles.xml').decode('utf-8')
        except KeyError:
            styles_xml = ''

    company = COMPANIES[company_name]
    income_code = '00發票'

    font = detect_main_font(xml + styles_xml)
    xml = update_header_parties(xml, company_name, font=font)
    xml = fill_bank_account(xml, company)
    xml = replace_signature_section(xml, company, font=font)
    xml = _fix_income_code_for_uspace(xml)

    validate_xml(xml)

    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zin, \
         zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                data = _strip_highlights(xml).encode('utf-8')
            elif item.filename == 'word/styles.xml':
                data = _strip_highlights(data.decode('utf-8')).encode('utf-8')
            elif item.filename.startswith('word/header'):
                hdr = data.decode('utf-8')
                hdr = update_header_income_code(hdr, income_code)
                data = hdr.encode('utf-8')
            zout.writestr(item, data)

    output_bytes = output.getvalue()

    # Append bank account image page if available for this company
    img_filename = BANK_IMAGES.get(company_name)
    if img_filename:
        img_path = os.path.join(PIC_DIR, img_filename)
        if os.path.exists(img_path):
            with open(img_path, 'rb') as f:
                img_data = f.read()
            output_bytes = _append_bank_image(output_bytes, img_data, img_filename)

    base = original_filename[:-5] if original_filename.lower().endswith('.docx') else original_filename
    output_filename = f'{base}(對悠勢合約).docx'

    return output_bytes, output_filename, []


def convert_contract(docx_bytes: bytes, original_filename: str, income_code: str = '') -> tuple:
    """
    Convert a 對業主 contract to the 對悠勢 format.
    - 7.7 業主版公版（內文有模式 A/B/C）→ docxtpl 渲染悠勢版 7.7 V2 公版
    - 非公版（增補合約等）→ 原文不動，只替換甲乙方
    Returns (output_bytes, output_filename, warnings).
    Raises ValueError for user-facing errors.
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        input_xml = z.read('word/document.xml').decode('utf-8')

    plain = get_plain_text(input_xml)

    company_name = detect_company(plain)
    if not company_name:
        # Log first 300 chars of plain text for diagnosis
        import sys
        print(f'[converter] detect_company failed. plain[:300]={plain[:300]!r}', file=sys.stderr)
        raise ValueError('無法識別合約中的公司（志昌／瀚昱／毅源），請確認上傳的是對業主合約')

    if is_v77_owner_contract(plain):
        return _convert_v77(docx_bytes, original_filename, company_name, plain)

    return _convert_party_swap(docx_bytes, original_filename, company_name)
